"""Create a real Word (.docx) document for Jarvis.

The companion to `create_pdf`: without a Word WRITER, "write my CV as a Word
doc" made the model call `write_file` with a `.docx` name and dump raw text --
bytes Word can't open. `create_docx` writes a genuine .docx (title + paragraphs)
via python-docx (pure-Python, offline).

Safety mirrors the other file tools: the path is REFUSED unless it stays inside
the user's home; content is bounded; the tool never raises. Unlike PDF, .docx is
UTF-8 XML, so the text keeps its real characters (accents, curly quotes) -- only
the RETURNED status line is forced to ASCII.
"""

from .find import _coerce
from .organize import _ascii, _first_str, _resolve_under_home
from .registry import tool

MAX_PATH_LEN = 400
MAX_CONTENT = 200000
MAX_TITLE = 200


def _import_docx():
    try:
        from docx import Document
        return Document
    except Exception:
        return None


@tool(
    "create_docx",
    "Create a real, openable Word document (.docx) and save it. Use this "
    "WHENEVER the user wants something written AS A WORD DOCUMENT -- a CV/resume, "
    "cover letter, report, or letter 'as a word doc' / 'as a docx' / 'in Word'. "
    "Do NOT use write_file for a .docx -- that produces a file Word can't open; "
    "create_docx writes a genuine document. Give path (e.g. 'Desktop/CV.docx'), "
    "content (the FULL text), and optionally title (a heading at the top). Only "
    "the user's own folders are allowed.",
    {
        "type": "object",
        "properties": {
            "path": {"type": "string",
                     "description": "Where to save it, e.g. 'Desktop/CV.docx'."},
            "content": {"type": "string",
                        "description": "The full text of the document."},
            "title": {"type": "string",
                      "description": "Optional heading shown at the top."},
        },
        "required": ["path", "content"],
    },
)
def create_docx(path: str = "", content: str = "", title: str = "", **extra) -> str:
    raw = _first_str(path, extra.get("file"), extra.get("filename"),
                     extra.get("dest"), extra.get("document"), extra.get("output"),
                     extra.get("name"))
    raw = _coerce(raw, MAX_PATH_LEN)
    if not raw:
        return "Error: tell me where to save the document, sir (e.g. Desktop/CV.docx)."

    body = _first_str(content, extra.get("text"), extra.get("body"),
                      extra.get("essay"), extra.get("message"))
    if not isinstance(body, str) or not body.strip():
        return "Error: tell me what to put in the document, sir."

    head = _first_str(title, extra.get("heading"), extra.get("header"))

    p, err = _resolve_under_home(raw)
    if p is None:
        return err or "Error: that document path isn't valid, sir."
    if p.suffix.lower() != ".docx":
        p = p.with_suffix(".docx")
    if p.is_dir():
        return f"Error: '{_ascii(p.name)}' is a folder, sir; give me a file name."

    Document = _import_docx()
    if Document is None:
        return ("Error: I can't make Word documents yet, sir -- the "
                "'python-docx' package isn't installed. Install it with: "
                "pip install python-docx")

    text = body[:MAX_CONTENT]
    heading = head[:MAX_TITLE] if head else ""
    try:
        doc = Document()
        if heading:
            doc.add_heading(heading, level=0)
        for line in text.split("\n"):
            doc.add_paragraph(line.rstrip())
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        paras = len(doc.paragraphs)
    except Exception as e:
        return f"Error: couldn't create that document, sir ({_ascii(str(e))})."

    return f"Created a Word document ({paras} paragraphs) at {_ascii(str(p))}."
